"""
增强版RAG系统 - Enhanced Retrieval Augmented Generation
包含以下高级功能：
1. 混合检索（语义 + 关键词BM25）
2. 查询扩展和重写
3. 重排序（Reranker）
4. 文档分块优化
5. 自我判断和思维链推理
"""
import os
from typing import List, Dict, Tuple
from pathlib import Path
import re
from dataclasses import dataclass

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.chat_models.tongyi import ChatTongyi
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.docstore.document import Document
from langchain_core.prompts import ChatPromptTemplate
import docx
import pandas as pd


@dataclass
class RetrievalResult:
    """检索结果"""
    content: str
    source: str
    score: float
    metadata: Dict


class EnhancedRAG:
    """增强版RAG系统"""

    def __init__(self, api_key: str, folder_path: str = "./ku"):
        self.api_key = api_key
        self.folder_path = folder_path
        self.vector_store = None
        self.bm25_retriever = None
        self.ensemble_retriever = None
        self.llm = None
        self.documents = []  # 保存原始文档用于BM25

        # 初始化LLM
        self.llm = ChatTongyi(
            model_name="qwen-plus",
            temperature=0.3,
            dashscope_api_key=api_key,
            streaming=False,
        )

    def load_documents(self, max_files: int = 15) -> List[Document]:
        """加载知识库文档（增强版，加载更多文件）

        改进：
        - 增加加载文件数量（5 -> 15）
        - 提取文档标题作为元数据
        - 过滤低质量内容
        """
        documents = []
        folder = Path(self.folder_path)

        if not folder.exists():
            print(f"⚠️  知识库文件夹不存在: {self.folder_path}")
            return documents

        print(f"📚 正在加载知识库...", end='', flush=True)

        # 优先加载的核心文件
        priority_keywords = [
            "交易原则", "战法", "情绪流", "打板", "超短",
            "龙头", "选股", "操作", "正道", "刺客"
        ]

        loaded_files = []
        success_count = 0

        # 遍历文件，按优先级加载
        all_files = list(folder.glob("*"))

        # 按优先级排序
        def get_priority(file_path):
            name = file_path.name.lower()
            for idx, keyword in enumerate(priority_keywords):
                if keyword.lower() in name:
                    return idx
            return 999

        all_files_sorted = sorted(all_files, key=get_priority)

        for file_path in all_files_sorted:
            if success_count >= max_files:
                break

            if not file_path.is_file():
                continue

            file_name = file_path.name

            # 跳过临时文件和缓存目录
            if file_name.startswith('.~') or file_name.startswith('~$') or file_name.startswith('.'):
                continue

            # 跳过Excel文件（太大）
            if file_path.suffix in ['.xlsx', '.xls']:
                continue

            # 只处理txt、docx和csv
            if file_path.suffix not in ['.txt', '.docx', '.csv']:
                continue

            try:
                content = ""
                title = file_name.replace(file_path.suffix, '')

                if file_path.suffix == '.txt':
                    content = self._read_txt_file(str(file_path))
                elif file_path.suffix == '.docx':
                    content, title = self._read_docx_file_with_title(str(file_path))
                elif file_path.suffix == '.csv':
                    content, title = self._read_csv_file(str(file_path))

                # 过滤低质量内容
                if not content or len(content.strip()) < 50:
                    continue

                # 创建Document对象，包含丰富的元数据
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": file_name,
                        "title": title,
                        "path": str(file_path),
                        "file_type": file_path.suffix,
                        "char_count": len(content)
                    }
                )

                documents.append(doc)
                loaded_files.append(file_path)
                success_count += 1
                print('.', end='', flush=True)

            except Exception as e:
                print(f"\n⚠️  读取文件失败 {file_name}: {e}")
                continue

        print(f" ✓ 加载了 {success_count} 个文件")
        return documents

    def _read_txt_file(self, file_path: str) -> str:
        """读取txt文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()

    def _read_docx_file_with_title(self, file_path: str) -> Tuple[str, str]:
        """读取docx文件，提取标题

        返回：(内容, 标题)
        """
        try:
            doc = docx.Document(file_path)
            content = []
            title = Path(file_path).stem

            # 提取标题（通常第一段是标题）
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 50:  # 标题通常较短
                    title = first_para

            # 提取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            return '\n'.join(content), title
        except Exception as e:
            return "", Path(file_path).stem

    def _read_csv_file(self, file_path: str) -> Tuple[str, str]:
        """读取CSV文件并转换为结构化文本（针对屠龙表优化）

        返回：(内容, 标题)
        """
        try:
            # 尝试不同的编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return "", Path(file_path).stem

            title = Path(file_path).stem

            # 特殊处理：屠龙表 - 短线模式
            if "短线模式" in title:
                content = self._format_short_term_mode_csv(df, title)
            # 特殊处理：屠龙表 - 题材周期系统
            elif "题材周期系统" in title:
                content = self._format_theme_cycle_csv(df, title)
            # 特殊处理：屠龙表 - 竞价表
            elif "竞价表" in title:
                content = self._format_bidding_table_csv(df, title)
            else:
                # 默认格式
                content = self._format_generic_csv(df, title)

            return content, title

        except Exception as e:
            print(f"\n⚠️  CSV读取失败 {file_path}: {e}")
            return "", Path(file_path).stem

    def _format_short_term_mode_csv(self, df: pd.DataFrame, title: str) -> str:
        """专门格式化短线模式CSV，保持表格结构"""
        content_parts = [
            f"# {title}",
            "",
            "## 短线模式体系结构",
            "本表格包含三大类短线模式：",
            "1. 一波流接力模式",
            "2. 趋势接力/龙头接力主升模式",
            "3. 纯趋势模式",
            "",
        ]

        # 按行情种类分组
        current_category = None

        for idx, row in df.iterrows():
            # 提取关键字段
            category = str(row.get('行情种类', '')).strip()
            mode_type = str(row.get('模式种类', '')).strip()
            name = str(row.get('名字', '')).strip()

            if not name or name == 'nan':
                continue

            # 检测新分类
            if category and category != 'nan' and category != current_category:
                current_category = category
                content_parts.append(f"\n## 【{current_category}】")
                content_parts.append("")

            # 模式标题
            content_parts.append(f"### 模式：{name}")

            # 关键信息
            fields_to_extract = [
                ('模式种类', '类型'),
                ('可做周期', '适用周期'),
                ('模式胜率高场景', '胜率高场景'),
                ('盈亏比', '盈亏比'),
                ('模式条件', '模式条件'),
                ('买点', '买点'),
                ('卖点', '卖点'),
                ('仓位', '仓位'),
                ('惩罚机制及心态管理', '纪律要求'),
            ]

            for col, label in fields_to_extract:
                value = str(row.get(col, '')).strip()
                # 完全跳过空值，不显示任何文字
                if value and value != 'nan' and len(value) > 0:
                    # 多行内容处理
                    if '\n' in value:
                        content_parts.append(f"**{label}**：")
                        for line in value.split('\n'):
                            if line.strip():
                                content_parts.append(f"  - {line.strip()}")
                    else:
                        content_parts.append(f"**{label}**：{value}")

            content_parts.append("")  # 空行分隔

        return "\n".join(content_parts)

    def _format_bidding_table_csv(self, df: pd.DataFrame, title: str) -> str:
        """专门格式化竞价表CSV（时间序列记录）- 简洁版"""
        content_parts = [
            f"# {title}",
            "",
            "竞价表是每日市场情绪的记录表，记录关键指标用于判断市场状态。",
            "",
        ]

        # 过滤掉空行和无效数据
        df_valid = df[df['日期'].notna()].copy()

        if len(df_valid) == 0:
            return "\n".join(content_parts)

        # 按日期排序（降序，最新的在前）
        try:
            df_valid['日期'] = pd.to_datetime(df_valid['日期'], errors='coerce')
            df_valid = df_valid.sort_values('日期', ascending=False)
        except:
            pass

        # 显示最近30天的数据
        recent_days = min(30, len(df_valid))
        df_recent = df_valid.head(recent_days)

        content_parts.append(f"最近{recent_days}天数据：")
        content_parts.append("")

        # 关键字段列表
        key_columns = ['节点', '高位', '中位', '小票亏效', '大票亏效', '一字', '断板', '是否开仓']

        # 简洁格式：每天一行，所有字段横排
        for _, row in df_recent.iterrows():
            date_str = str(row.get('日期', ''))[:10]
            if not date_str or date_str == 'NaT':
                continue

            # 收集该日期的所有字段值
            fields = []
            for col in key_columns:
                value = str(row.get(col, '')).strip()
                if value and value != 'nan' and value != '':
                    fields.append(f"{col}={value}")

            # 一行展示
            if fields:
                content_parts.append(f"[{date_str}] " + " | ".join(fields))

        content_parts.append("")
        return "\n".join(content_parts)

    def _format_theme_cycle_csv(self, df: pd.DataFrame, title: str) -> str:
        """专门格式化题材周期系统CSV"""
        content_parts = [
            f"# {title}",
            "",
            "## 题材周期系统说明",
            "题材在市场中经历不同的生命周期阶段，每个阶段有不同的表现特征和操作机会。",
            "",
        ]

        # 获取列名（题材节点）
        # 第一行是：可参与题材, 题材节点, 发酵, 强更强, 首次分歧, 弱转强, 高潮, 2次分歧加入轮动
        # 跳过第一列"可参与题材"
        if len(df) == 0:
            return "\n".join(content_parts)

        # 提取题材节点（列名）
        col_names = df.columns.tolist()
        # col_names[0] 是 "可参与题材"，col_names[1:] 是各个阶段
        stages = []
        for i in range(1, len(col_names)):
            col_name = col_names[i]
            # 从第一行（题材节点行）获取阶段名称
            if len(df) > 0:
                stage_name = str(df.iloc[0, i]).strip()
                if stage_name and stage_name != 'nan':
                    stages.append((i, stage_name))

        # 如果没有找到阶段，使用列索引
        if not stages:
            stages = [(i, col_names[i]) for i in range(1, len(col_names))]

        # 按阶段格式化
        for col_idx, stage_name in stages:
            content_parts.append(f"### 【{stage_name}】")
            content_parts.append("")

            # 遍历每一行（维度）
            for row_idx in range(1, len(df)):  # 跳过第0行（题材节点行）
                dimension_name = str(df.iloc[row_idx, 0]).strip()  # 第一列是维度名
                if not dimension_name or dimension_name == 'nan':
                    continue

                value = str(df.iloc[row_idx, col_idx]).strip()
                if value and value != 'nan' and value != '无':
                    # 多行内容处理
                    if '\n' in value:
                        content_parts.append(f"**{dimension_name}**：")
                        for line in value.split('\n'):
                            if line.strip():
                                content_parts.append(f"  - {line.strip()}")
                    else:
                        content_parts.append(f"**{dimension_name}**：{value}")

            content_parts.append("")  # 阶段间空行

        return "\n".join(content_parts)

    def _format_generic_csv(self, df: pd.DataFrame, title: str) -> str:
        """通用CSV格式化方法"""
        content_parts = [f"# {title}\n"]

        for idx, row in df.iterrows():
            row_text_parts = []
            for col_name, value in row.items():
                if pd.notna(value) and str(value).strip():
                    row_text_parts.append(f"**{col_name}**: {value}")

            if row_text_parts:
                content_parts.append(f"## 条目 {idx + 1}")
                content_parts.append("\n".join(row_text_parts))
                content_parts.append("")

        content = "\n".join(content_parts)

        # 如果内容太短，使用表格形式
        if len(content) < 100:
            content = f"# {title}\n\n{df.to_string(index=False)}"

        return content

    def build_enhanced_index(self):
        """构建增强版索引

        改进：
        1. 优化文档分块（更大的chunk，更好的overlap）
        2. 创建向量索引（FAISS）
        3. 创建BM25关键词索引
        4. 组合成混合检索器
        5. 并发加载所有文档，保持性能
        """
        # 加载所有文档（使用None表示不限制数量）
        self.documents = self.load_documents(max_files=None)

        if not self.documents:
            print("⚠️  没有加载到知识文档")
            return False

        # 优化的文本分块策略
        print("📝 正在分块文档...", end='', flush=True)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # 增大chunk以保持上下文完整性
            chunk_overlap=200,  # 增加overlap避免信息断裂
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
            keep_separator=True
        )

        split_docs = text_splitter.split_documents(self.documents)

        # 为每个chunk添加额外元数据
        for i, doc in enumerate(split_docs):
            doc.metadata['chunk_id'] = i
            doc.metadata['chunk_size'] = len(doc.page_content)

            # 提取关键词（简单版：提取前100字符作为摘要）
            doc.metadata['summary'] = doc.page_content[:100] + "..."

        print(f" ✓ 生成 {len(split_docs)} 个文本块")

        # 1. 构建向量索引（可选，失败时降级到纯BM25）
        print("🔍 正在构建向量索引...", end='', flush=True)
        vector_retriever = None
        try:
            embeddings = DashScopeEmbeddings(
                model="text-embedding-v1",
                dashscope_api_key=self.api_key
            )
            self.vector_store = FAISS.from_documents(
                documents=split_docs,
                embedding=embeddings
            )
            vector_retriever = self.vector_store.as_retriever(
                search_kwargs={"k": 4}
            )
            print(" ✓")
        except Exception as e:
            print(f"\n⚠️  向量索引构建失败: {e}")
            print("   将降级使用BM25关键词检索")

        # 2. 构建BM25关键词索引（必须成功）
        print("🔍 正在构建BM25索引...", end='', flush=True)
        try:
            self.bm25_retriever = BM25Retriever.from_documents(split_docs)
            self.bm25_retriever.k = 8  # 增加返回结果数
            print(" ✓")
        except Exception as e:
            print(f"\n❌ BM25索引构建失败: {e}")
            return False

        # 3. 组合混合检索器
        print("🔍 正在构建检索器...", end='', flush=True)
        if vector_retriever:
            # 向量+BM25混合检索
            self.ensemble_retriever = EnsembleRetriever(
                retrievers=[vector_retriever, self.bm25_retriever],
                weights=[0.6, 0.4]
            )
            print(" ✓ (混合检索)")
        else:
            # 仅BM25检索
            self.ensemble_retriever = self.bm25_retriever
            print(" ✓ (BM25检索)")

        print("✅ 增强版RAG系统初始化完成！")
        return True

    def expand_query(self, query: str) -> List[str]:
        """查询扩展 - 使用LLM生成相关查询变体

        例如："如何判断龙头股" ->
        ["如何判断龙头股", "龙头股的特征", "龙头股选股方法", "如何找到龙头"]
        """
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "你是一个查询扩展专家。给定一个用户查询，生成3个相关但措辞不同的查询变体。\n"
             "这些变体应该帮助检索到更多相关信息。\n"
             "每个查询一行，不要编号，不要添加任何解释。"),
            ("human", "原始查询：{query}\n\n生成3个查询变体：")
        ])

        try:
            response = self.llm.invoke(prompt.format_messages(query=query))
            expanded = response.content.strip().split('\n')

            # 清理结果
            expanded = [q.strip() for q in expanded if q.strip()]
            expanded = [re.sub(r'^\d+[\.\、]\s*', '', q) for q in expanded]  # 移除编号

            # 添加原始查询
            result = [query] + expanded[:3]
            return result
        except:
            # 扩展失败，返回原始查询
            return [query]

    def retrieve_with_rerank(self, query: str, top_k: int = 5) -> List[RetrievalResult]:
        """混合检索 + 重排序

        流程：
        1. 查询扩展（生成多个查询变体）
        2. 混合检索（向量 + BM25）
        3. 重排序（使用LLM对结果打分）
        4. 返回top_k结果
        """
        # 1. 查询扩展
        print(f"🔍 正在扩展查询...", end='', flush=True)
        expanded_queries = self.expand_query(query)
        print(f" ✓ 生成 {len(expanded_queries)} 个查询")

        # 2. 对每个扩展查询进行混合检索
        all_docs = []
        seen_content = set()  # 去重

        for q in expanded_queries:
            try:
                docs = self.ensemble_retriever.get_relevant_documents(q)

                for doc in docs:
                    # 去重（基于内容的前100字符）
                    content_hash = doc.page_content[:100]
                    if content_hash not in seen_content:
                        seen_content.add(content_hash)
                        all_docs.append(doc)
            except Exception as e:
                print(f"\n⚠️  检索失败: {e}")
                continue

        if not all_docs:
            return []

        print(f"📚 检索到 {len(all_docs)} 个候选文档")

        # 3. 重排序（使用LLM对相关性打分）
        print(f"🎯 正在重排序...", end='', flush=True)
        reranked = self._rerank_documents(query, all_docs)
        print(" ✓")

        # 返回top_k结果
        return reranked[:top_k]

    def _rerank_documents(self, query: str, documents: List[Document]) -> List[RetrievalResult]:
        """重排序 - 使用LLM评估文档与查询的相关性

        为每个文档打分（0-10），然后按分数排序
        """
        results = []

        # 批量处理以提高效率
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "你是一个文档相关性评分专家。给定一个查询和一段文档内容，"
             "评估文档对回答查询的相关性。\n"
             "只输出一个0-10的数字，10表示非常相关，0表示完全不相关。\n"
             "不要输出任何解释，只输出数字。"),
            ("human",
             "查询：{query}\n\n"
             "文档内容：{content}\n\n"
             "相关性分数（0-10）：")
        ])

        for doc in documents:
            try:
                # 截取文档前500字符用于评分（节省token）
                content_preview = doc.page_content[:500]

                response = self.llm.invoke(
                    prompt.format_messages(query=query, content=content_preview)
                )

                # 提取分数
                score_text = response.content.strip()
                score = float(re.search(r'\d+\.?\d*', score_text).group())
                score = max(0, min(10, score))  # 限制在0-10

            except:
                # 打分失败，给默认分数
                score = 5.0

            result = RetrievalResult(
                content=doc.page_content,
                source=doc.metadata.get('source', '未知'),
                score=score,
                metadata=doc.metadata
            )
            results.append(result)

        # 按分数降序排序
        results.sort(key=lambda x: x.score, reverse=True)

        return results

    def retrieve(self, query: str, top_k: int = 5, use_rerank: bool = True) -> str:
        """检索知识

        参数：
        - query: 查询问题
        - top_k: 返回结果数量
        - use_rerank: 是否使用重排序（更慢但更准确）

        返回：格式化的检索结果文本
        """
        if not self.ensemble_retriever:
            return "知识库未初始化"

        try:
            if use_rerank:
                results = self.retrieve_with_rerank(query, top_k=top_k)
            else:
                # 简单检索（不重排序）
                docs = self.ensemble_retriever.get_relevant_documents(query)
                results = [
                    RetrievalResult(
                        content=doc.page_content,
                        source=doc.metadata.get('source', '未知'),
                        score=8.0,
                        metadata=doc.metadata
                    )
                    for doc in docs[:top_k]
                ]

            if not results:
                return "未找到相关知识"

            # 格式化输出
            output = []
            for i, result in enumerate(results, 1):
                title = result.metadata.get('title', result.source)
                output.append(
                    f"【知识 {i}】来源：{title} (相关性: {result.score:.1f}/10)\n"
                    f"{result.content}\n"
                )

            return "\n---\n".join(output)

        except Exception as e:
            return f"检索失败: {str(e)}"

    def answer_with_reasoning(self, query: str, context: str = "") -> str:
        """带思维链的回答生成 - 增强自我判断能力

        改进：
        1. 先检索相关知识
        2. 让模型进行思维链推理
        3. 生成最终答案
        """
        # 1. 检索知识
        print("🔍 正在检索知识库...")
        knowledge = self.retrieve(query, top_k=5, use_rerank=True)

        # 2. 带思维链的推理prompt
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "你是一位专业的股票交易分析专家。请基于提供的知识库内容和实时信息，"
             "运用思维链推理来回答问题。\n\n"
             "【思维链推理步骤】\n"
             "1. **理解问题**：明确用户在问什么\n"
             "2. **分析知识**：从知识库中找到相关要点\n"
             "3. **结合实际**：如果有实时数据，结合分析\n"
             "4. **形成判断**：基于知识和数据，给出专业判断\n"
             "5. **给出建议**：提供可操作的建议\n\n"
             "【回答格式】\n"
             "💭 **思考过程**：（简要说明你的分析思路）\n"
             "📊 **核心观点**：（3-5个关键点）\n"
             "✅ **操作建议**：（具体可行的建议）\n\n"
             "【注意】\n"
             "- 必须基于知识库内容，不要编造\n"
             "- 如果知识库中没有直接答案，说明需要结合实时数据判断\n"
             "- 保持客观专业，不做过度承诺\n"
             "- 如果不确定，明确指出不确定性"),
            ("human",
             "【用户问题】\n{query}\n\n"
             "【知识库检索结果】\n{knowledge}\n\n"
             "【实时上下文】\n{context}\n\n"
             "请运用思维链推理回答：")
        ])

        try:
            response = self.llm.invoke(
                prompt.format_messages(
                    query=query,
                    knowledge=knowledge,
                    context=context if context else "（暂无实时数据）"
                )
            )

            return response.content

        except Exception as e:
            return f"回答生成失败: {str(e)}"


def test_enhanced_rag():
    """测试增强版RAG"""
    print("=" * 70)
    print("增强版RAG系统 - 测试")
    print("=" * 70)

    api_key = os.environ.get("DASHSCOPE_API_KEY")
    if not api_key:
        print("\n⚠️  警告：未检测到 DASHSCOPE_API_KEY 环境变量")
        return

    # 初始化
    rag = EnhancedRAG(api_key=api_key, folder_path="./ku")

    # 构建索引
    if not rag.build_enhanced_index():
        print("❌ 索引构建失败")
        return

    # 测试问题
    test_questions = [
        "如何判断市场情绪周期？",
        "打板战法的核心要点是什么？",
        "龙头股有什么特征？",
    ]

    for i, question in enumerate(test_questions, 1):
        print(f"\n{'=' * 70}")
        print(f"📝 测试 {i}/{len(test_questions)}: {question}")
        print(f"{'=' * 70}\n")

        # 测试检索
        print("【检索测试】")
        result = rag.retrieve(question, top_k=3, use_rerank=True)
        print(result)
        print()

        # 测试推理回答
        print("\n【推理回答测试】")
        answer = rag.answer_with_reasoning(question)
        print(answer)

        import time
        time.sleep(2)

    print(f"\n{'=' * 70}")
    print("✓ 测试完成")
    print(f"{'=' * 70}")


if __name__ == "__main__":
    test_enhanced_rag()
